from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCX_PATH = Path(r"C:\Users\ACER-A315-59\OneDrive\Desktop\Guia_Exposicion_IA.docx")
BACKUP_PATH = Path(r"C:\Users\ACER-A315-59\OneDrive\Desktop\Proyecto final IA\Guia_Exposicion_IA_backup_before_fases_4_5.docx")


ACCENT = "1F4E79"
LIGHT = "EAF2F8"
HEADER = "D9EAF7"
SOFT = "F6F8FA"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_table(table, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    if widths:
        set_col_widths(table, widths)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
            if r_idx == 0:
                set_cell_shading(cell, HEADER)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def set_table_rows(table, rows):
    for row_idx, row_data in enumerate(rows):
        cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(row_data):
            cells[col_idx].text = str(value)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Arial"
    for run in p.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.color.rgb = RGBColor(31, 78, 121)
            run.font.size = Pt(16)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)
        else:
            run.font.color.rgb = RGBColor(44, 62, 80)
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(5)
    return p


def add_para(doc, text="", bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        text = text if text.startswith(" ") else " " + text
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        r = p.add_run("- " + item)
        r.font.name = "Arial"
        r.font.size = Pt(10)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title + " ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(31, 78, 121)
    r = p.add_run(body)
    r.font.name = "Arial"
    r.font.size = Pt(10)
    doc.add_paragraph()


def delete_from_phase_4(doc):
    body = doc._body._element
    children = list(body)
    start = None
    for idx, child in enumerate(children):
        texts = child.xpath(".//w:t/text()")
        if any("FASE 4" in t for t in texts):
            start = idx
            break
    if start is None:
        raise RuntimeError("No se encontro el marcador FASE 4 en el documento.")
    for child in children[start:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def add_phase_4(doc):
    add_heading(doc, "FASE 4  -  Modelado y entrenamiento", 1)
    add_para(
        doc,
        "En esta fase se entrenaron cinco algoritmos supervisados sobre la matriz TF-IDF generada en la Fase 3: Regresion Logistica, Random Forest, Arbol de Decision, Red Neuronal MLP y una DNN implementada con MLPClassifier de mayor profundidad. Todos se evaluaron bajo una misma particion 80/20, con escalado MaxAbsScaler y validacion cruzada estratificada.",
    )
    add_callout(
        doc,
        "Idea central para defender:",
        "la fase no buscaba probar un solo modelo, sino comparar familias distintas bajo condiciones controladas para identificar cual generaliza mejor en un problema multiclase de 30 categorias.",
    )

    add_heading(doc, "4.1  Que se hizo y por que", 2)
    table = doc.add_table(rows=6, cols=4)
    set_table_rows(
        table,
        [
            ["Decision", "Que se hizo", "Por que se hizo", "Impacto esperado"],
            ["Validacion cruzada", "GridSearchCV con StratifiedKFold.", "Evita depender de una unica particion y conserva representacion de las 30 clases en cada fold.", "Seleccion de hiperparametros mas estable y comparable."],
            ["Metrica de busqueda", "Accuracy en validacion cruzada.", "El dataset quedo balanceado a 3,000 registros por clase; por eso el accuracy deja de estar dominado por clases grandes.", "Optimiza rendimiento global sin favorecer una clase mayoritaria."],
            ["Escalado", "MaxAbsScaler ajustado solo sobre train.", "Mantiene la matriz dispersa TF-IDF y evita fuga de informacion del test.", "Mejora compatibilidad con modelos lineales y redes sin densificar la matriz."],
            ["Reproducibilidad", "random_state=42 y particion estratificada.", "Permite repetir resultados y sostener la defensa frente al jurado.", "Resultados trazables y verificables."],
            ["Comparacion plural", "Modelos lineales, arboles, ensemble y redes.", "Cada familia aprende patrones diferentes: pesos lexicos, reglas, votaciones y relaciones no lineales.", "Permite justificar la eleccion final con evidencia, no con preferencia."],
        ],
    )
    style_table(table, [1.25, 2.05, 2.75, 2.25])

    add_heading(doc, "4.2  Modelos entrenados y justificacion tecnica", 2)
    table = doc.add_table(rows=6, cols=5)
    set_table_rows(
        table,
        [
            ["Modelo", "Hiperparametros clave", "Justificacion tecnica", "Resultado", "Lectura del impacto"],
            ["Regresion Logistica", "C=5, penalty=L2, class_weight='balanced', solver lbfgs.", "Es fuerte en texto con TF-IDF porque aprende pesos directos por termino y suele generalizar bien en espacios dispersos de alta dimension.", "CV 0.6621; Test 0.6788; F1 macro 0.6775.", "Fue el mejor modelo: buen equilibrio entre precision, interpretabilidad y costo computacional."],
            ["Random Forest", "200 arboles, max_depth=None, max_features='sqrt', min_samples_leaf=1.", "Evalua reglas no lineales y reduce varianza frente a un arbol individual mediante votacion.", "CV 0.6390; Test 0.6563; F1 macro 0.6494.", "Competitivo, pero menos eficaz en TF-IDF disperso y mas costoso que el modelo lineal."],
            ["Arbol de Decision", "max_depth=None, criterion='gini', sin class_weight.", "Sirve como modelo interpretable de reglas y linea base no lineal sencilla.", "CV 0.5712; Test 0.5958; F1 macro 0.5895.", "Fue el mas debil; confirma que un solo arbol sobreajusta y no captura bien la complejidad textual."],
            ["Red Neuronal MLP", "hidden_layer_sizes=(100,), alpha=0.001, learning_rate='adaptive'.", "Busca relaciones no lineales entre terminos TF-IDF con early stopping para controlar sobreajuste.", "CV 0.6569; Test 0.6711; F1 macro 0.6702.", "Quedo muy cerca de Regresion Logistica, pero con mayor tiempo de entrenamiento."],
            ["DNN", "hidden_layer_sizes=(256,128), alpha=0.0001, learning_rate_init=0.001.", "Prueba una arquitectura mas profunda para capturar combinaciones jerarquicas de patrones lexicos.", "CV 0.6281; Test 0.6554; F1 macro 0.6538.", "No supero al MLP ni al modelo lineal; la profundidad aumento costo sin mejorar generalizacion."],
        ],
    )
    style_table(table, [1.15, 1.8, 2.4, 1.45, 1.9])

    add_heading(doc, "4.3  Justificacion tecnica de las decisiones principales", 2)
    add_bullets(
        doc,
        [
            "Se uso GridSearchCV porque la seleccion manual de hiperparametros puede sesgar la comparacion. La busqueda permite evaluar combinaciones bajo la misma regla objetiva.",
            "Se uso validacion cruzada estratificada porque el problema tiene 30 clases; sin estratificacion algunos folds podrian quedar pobres en ciertas categorias.",
            "Se mantuvo TF-IDF como representacion porque los modelos seleccionados trabajan con vectores numericos y el vocabulario discriminativo es una senal fuerte para noticias.",
            "Se escalo con MaxAbsScaler porque StandardScaler puede densificar matrices dispersas; eso elevaria memoria y tiempo de ejecucion sin necesidad.",
            "Se incluyo F1 macro en la evaluacion final porque mide el equilibrio por clase y no solo el porcentaje global de aciertos.",
        ],
    )

    add_heading(doc, "4.4  Impacto en el resultado", 2)
    add_para(
        doc,
        "La fase de modelado demostro que el pipeline de preprocesamiento era solido: incluso modelos muy distintos quedaron claramente por encima del azar. En un problema de 30 clases, una prediccion aleatoria tendria cerca de 3.33% de acierto; el mejor modelo alcanzo 67.88%. La Regresion Logistica gano porque TF-IDF produce un espacio de alta dimension donde los limites lineales suelen ser muy competitivos.",
    )

    add_heading(doc, "4.5  Guion de exposicion oral", 2)
    add_para(
        doc,
        "En la fase cuatro pasamos de tener texto preparado a entrenar modelos supervisados. Para que la comparacion fuera justa, usamos la misma matriz TF-IDF, la misma particion train/test y la misma estrategia de validacion cruzada estratificada. Entrenamos cinco enfoques: Regresion Logistica, Arbol de Decision, Random Forest, MLP y DNN. La razon de probar varios modelos fue cubrir familias distintas: un modelo lineal, modelos basados en arboles, un ensemble y redes neuronales.",
    )
    add_para(
        doc,
        "La Regresion Logistica funciono especialmente bien porque los datos de texto vectorizados con TF-IDF suelen separarse mediante pesos lexicos: algunas palabras o bigramas son muy informativos para categorias como tecnologia, religion, deporte o crimen. Las redes neuronales tambien fueron competitivas, pero con mayor costo computacional. En cambio, el arbol individual fue el mas bajo, lo cual era esperable porque puede sobreajustar y le cuesta generalizar en espacios dispersos con muchas caracteristicas.",
    )
    add_para(
        doc,
        "La decision tecnica mas importante fue no elegir el modelo por intuicion, sino por evidencia: todos pasaron por busqueda de hiperparametros y luego se evaluaron con datos de prueba no vistos.",
    )

    add_heading(doc, "4.6  Preguntas probables del jurado y respuestas", 2)
    table = doc.add_table(rows=7, cols=2)
    set_table_rows(
        table,
        [
            ["Pregunta probable", "Respuesta sugerida"],
            ["Por que usaron accuracy para GridSearch si el dataset original estaba desbalanceado?", "Porque antes del modelado se balanceo el dataset a 3,000 registros por clase. En ese contexto, accuracy es valida para optimizar rendimiento global. Aun asi, en la evaluacion final reportamos F1 macro para comprobar equilibrio por clase."],
            ["Por que Regresion Logistica supero a redes neuronales?", "Porque TF-IDF genera variables lexicas muy discriminativas y de alta dimension. En ese escenario, un clasificador lineal regularizado suele generalizar muy bien. Las redes agregaron complejidad, pero no suficiente mejora para compensar el costo."],
            ["Por que escalaron datos TF-IDF?", "Para estabilizar modelos sensibles a escala, especialmente redes y regresion. Se uso MaxAbsScaler porque conserva la naturaleza dispersa de la matriz y se ajusto solo con train para evitar data leakage."],
            ["Por que incluir un Arbol de Decision si rindio bajo?", "Porque sirve como linea base interpretable no lineal. Su bajo resultado aporta evidencia: un solo arbol no es suficiente para este tipo de representacion textual."],
            ["Como controlaron el sobreajuste?", "Con validacion cruzada, regularizacion L2/alpha, early stopping en redes, evaluacion final en test no visto y comparacion entre CV y test."],
            ["Que limitacion tuvo el entrenamiento?", "El costo computacional fue alto, especialmente en DNN. Ademas, el oversampling puede repetir textos minoritarios y aumentar riesgo de memorizar ejemplos."],
        ],
    )
    style_table(table, [2.55, 5.55])

    add_heading(doc, "4.7  Observaciones para fortalecer la defensa", 2)
    add_bullets(
        doc,
        [
            "Reconocer que el oversampling por repeticion ayuda al balance, pero puede aumentar sobreajuste; defenderlo como una decision necesaria y controlada con test separado.",
            "Aclarar que el mejor modelo no fue el mas complejo, sino el que generalizo mejor. Esa es una conclusion valiosa, no una debilidad.",
            "Mencionar que una mejora futura seria usar embeddings o transformers en lugar de TF-IDF, pero con mayor costo computacional.",
            "Si preguntan por interpretabilidad, destacar que Regresion Logistica permite analizar pesos por palabra y explicar por que una noticia cae en cierta categoria.",
        ],
    )


def add_phase_5(doc):
    add_heading(doc, "FASE 5  -  Evaluacion, comparacion y conclusiones", 1)
    add_para(
        doc,
        "La fase final compara los modelos con metricas cuantitativas y extrae conclusiones defendibles. La evaluacion se hizo con el conjunto de prueba, separado desde la Fase 3, para medir generalizacion sobre noticias no vistas durante el entrenamiento.",
    )

    add_heading(doc, "5.1  Que se evaluo y por que", 2)
    table = doc.add_table(rows=5, cols=4)
    set_table_rows(
        table,
        [
            ["Metrica", "Que mide", "Por que se uso", "Impacto en la decision"],
            ["Accuracy", "Proporcion total de aciertos.", "Es interpretable y valida porque el test quedo balanceado por clase.", "Permite ordenar modelos por desempeno global."],
            ["F1 macro", "Promedio del F1 de cada clase con igual peso.", "Evita que categorias faciles oculten bajo rendimiento en categorias dificiles.", "Confirma si el modelo es equilibrado en las 30 clases."],
            ["Classification report", "Precision, recall y F1 por categoria.", "Permite detectar clases fuertes y debiles.", "Ayuda a explicar errores y posibles mejoras."],
            ["Matriz de confusion", "Donde se equivoca cada modelo.", "Muestra confusiones entre categorias semanticamente cercanas.", "Convierte el error en informacion accionable."],
        ],
    )
    style_table(table, [1.35, 2.15, 2.35, 2.15])

    add_heading(doc, "5.2  Comparacion final de modelos", 2)
    table = doc.add_table(rows=6, cols=5)
    set_table_rows(
        table,
        [
            ["Ranking", "Modelo", "Accuracy CV", "Accuracy Test", "F1 macro Test"],
            ["1", "Regresion Logistica", "0.6621", "0.6788", "0.6775"],
            ["2", "Red Neuronal MLP", "0.6569", "0.6711", "0.6702"],
            ["3", "Random Forest", "0.6390", "0.6563", "0.6494"],
            ["4", "DNN", "0.6281", "0.6554", "0.6538"],
            ["5", "Arbol de Decision", "0.5712", "0.5958", "0.5895"],
        ],
    )
    style_table(table, [0.65, 2.4, 1.25, 1.25, 1.25])

    add_heading(doc, "5.3  Interpretacion de resultados", 2)
    add_bullets(
        doc,
        [
            "La Regresion Logistica fue el modelo ganador con 67.88% de accuracy y 67.75% de F1 macro en test.",
            "El MLP quedo muy cerca del ganador, con 67.11% de accuracy, pero exigio mas tiempo de entrenamiento.",
            "Random Forest y DNN fueron competitivos, aunque no superaron al modelo lineal. Esto indica que mayor complejidad no garantizo mejor generalizacion.",
            "El Arbol de Decision fue el mas debil, lo que confirma que una estructura unica de reglas no basta para representar bien 30 categorias de texto.",
            "Las clases con vocabulario mas distintivo, como GOOD NEWS, LATINO VOICES, SCIENCE, TECH, RELIGION o U.S. NEWS, presentaron mejores F1; las mas ambiguas, como BUSINESS, IMPACT, POLITICS, ENTERTAINMENT o PARENTING, concentraron errores.",
        ],
    )

    add_heading(doc, "5.4  Conclusiones defendibles", 2)
    add_para(
        doc,
        "El objetivo se cumplio: se construyo un sistema capaz de clasificar automaticamente noticias en 30 categorias usando NLP y Machine Learning. El resultado no debe leerse como un simple 67.88%, sino como un desempeno solido para un problema multicategoria amplio, donde el azar estaria alrededor de 3.33%.",
    )
    add_para(
        doc,
        "La principal conclusion tecnica es que la calidad del pipeline de datos fue mas determinante que la complejidad del algoritmo. Despues de limpiar, fusionar categorias similares, balancear y vectorizar con TF-IDF, un modelo lineal regularizado supero a alternativas mas complejas.",
    )
    add_para(
        doc,
        "La comparacion tambien muestra que las confusiones no son aleatorias: aparecen principalmente entre categorias semanticamente cercanas. Esto indica una limitacion natural del dataset y de los titulares cortos, no necesariamente una falla del modelo.",
    )

    add_heading(doc, "5.5  Guion de exposicion oral", 2)
    add_para(
        doc,
        "En la fase cinco evaluamos los modelos con datos de prueba que no fueron usados durante el entrenamiento. Usamos accuracy para medir el acierto global y F1 macro para revisar si el modelo se comportaba bien en todas las categorias. Esta segunda metrica era importante porque nuestro problema tiene 30 clases y no queriamos que el resultado dependiera solo de categorias faciles.",
    )
    add_para(
        doc,
        "El mejor resultado fue Regresion Logistica, con 67.88% de accuracy y 67.75% de F1 macro. El segundo fue MLP, con 67.11%, muy cercano, pero con mayor costo computacional. Random Forest y DNN quedaron alrededor del 65%, y el Arbol de Decision fue el mas bajo, con 59.58%. Esto nos deja una conclusion clara: para este dataset y esta representacion TF-IDF, el modelo mas complejo no fue el mejor; el mejor fue el que equilibro generalizacion, estabilidad e interpretabilidad.",
    )
    add_para(
        doc,
        "Tambien analizamos las clases donde el modelo acierta mejor y donde se confunde. Las categorias con vocabulario propio, como tecnologia, ciencia, religion o noticias positivas, se clasifican mejor. Las categorias mas amplias o cercanas semanticamente, como politica, negocios, impacto o entretenimiento, generan mas errores. Por eso nuestras conclusiones no solo reportan numeros, sino que explican el comportamiento del sistema.",
    )

    add_heading(doc, "5.6  Preguntas probables del jurado y respuestas", 2)
    table = doc.add_table(rows=8, cols=2)
    set_table_rows(
        table,
        [
            ["Pregunta probable", "Respuesta sugerida"],
            ["Un 67.88% es suficiente?", "Para 30 categorias es un resultado solido. El azar seria 3.33%, y ademas varias categorias son semanticamente cercanas. No se presenta como perfecto, sino como un clasificador funcional y mejorable."],
            ["Por que no gano la DNN?", "Porque la representacion TF-IDF ya entrega senales lexicas fuertes. La DNN agrego capacidad, pero tambien costo y riesgo de sobreajuste. En estos datos, la frontera lineal regularizada fue mas efectiva."],
            ["Que modelo elegirian para produccion?", "Regresion Logistica, porque fue el mejor en test, es rapida, reproducible e interpretable. Si el objetivo fuera maxima exactitud con mas recursos, se explorarian transformers."],
            ["Que categorias son mas problematicas?", "Las de frontera semantica amplia o vocabulario compartido: POLITICS, BUSINESS, IMPACT, ENTERTAINMENT y PARENTING. No son errores aislados; responden al solapamiento tematico del corpus."],
            ["Como saben que no hubo data leakage?", "El escalador se ajusto solo con train, la evaluacion final se hizo en test separado y la seleccion de hiperparametros se hizo dentro de la validacion cruzada del entrenamiento."],
            ["Que cambiarian con mas tiempo?", "Probar embeddings contextuales, calibrar probabilidades, analizar importancia de terminos, revisar manualmente clases ambiguas y usar validacion externa con noticias recientes."],
            ["Por que usar F1 macro?", "Porque da el mismo peso a cada categoria. En un problema multiclase, esto evita que el desempeno global oculte clases con bajo recall o baja precision."],
        ],
    )
    style_table(table, [2.55, 5.55])

    add_heading(doc, "5.7  Observaciones para fortalecer la defensa", 2)
    add_bullets(
        doc,
        [
            "No vender el modelo como perfecto. Defenderlo como un sistema solido, medido con criterios justos y consciente de sus limites.",
            "Llevar claro el argumento del azar: 1 de 30 clases equivale a 3.33%, por eso 67.88% es significativo.",
            "Explicar que F1 macro respalda el resultado porque muestra equilibrio, no solo aciertos acumulados.",
            "Si cuestionan el balanceo, reconocer que repetir ejemplos minoritarios es una limitacion; la respuesta es que se mitigo con test separado y validacion cruzada.",
            "Si preguntan por el futuro, proponer transformers como BERT o RoBERTa, pero aclarar que requieren mas recursos y una evaluacion comparativa nueva.",
            "Recordar que la fusion de categorias similares fue clave: sin ella, el modelo tendria que separar clases casi indistinguibles por texto.",
        ],
    )


def main():
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(str(DOCX_PATH))
    delete_from_phase_4(doc)
    add_phase_4(doc)
    doc.add_page_break()
    add_phase_5(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    main()
